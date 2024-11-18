import re
from docx import Document

# Specifica i file di input e output
input_file = "/Users/boss/Desktop/Testi/SignoreDellaNave.docx"
output_file = "/Users/boss/Desktop/Testi/file_modificato.docx"

# Carica il documento esistente
doc = Document(input_file)

# Crea un nuovo documento per le modifiche
new_doc = Document()

# Itera su tutti i paragrafi del file
for paragraph in doc.paragraphs:
    text = paragraph.text
    if text.strip():  # Se il paragrafo non è vuoto
        # Usa regex per gestire i casi specifici
        modified_text = re.sub(r'\-', r'\n-', text)  # Aggiungi \n- dopo i trattini
        modified_text = re.sub(r'\.', r'.\n', modified_text)  # Aggiungi .\n dopo i punti
        
        for line in modified_text.split('\n'):
            new_doc.add_paragraph(line.strip())  # Aggiungi ogni riga come nuovo paragrafo
    else:
        new_doc.add_paragraph("")  # Mantieni i paragrafi vuoti

# Salva il documento modificato
new_doc.save(output_file)
print(f"File modificato salvato come '{output_file}'.")